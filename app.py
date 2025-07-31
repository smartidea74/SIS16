import streamlit as st
from docx import Document
from datetime import date
import io
import math

st.set_page_config(page_title="Сметка за изплатени суми", layout="wide")

st.markdown("""
    <style>
    .stTabs [data-baseweb="tab-list"] {
        border-bottom: 2px solid #ccc;
    }
    .stTabs [data-baseweb="tab"] {
        font-size: 16px;
        padding: 10px 20px;
        margin-right: 8px;
        background-color: #f0f2f6;
        border-radius: 5px 5px 0 0;
        border: 1px solid #ccc;
        border-bottom: none;
    }
    .stTabs [aria-selected="true"] {
        background-color: white;
        color: black;
        font-weight: bold;
        border-bottom: 2px solid white;
    }
    .result-box {
        background-color: #f9f9f9;
        border-left: 5px solid #4CAF50;
        padding: 10px;
        margin-top: 20px;
    }
    </style>
""", unsafe_allow_html=True)

MAX_INSURANCE_INCOME = 4130.00

def calculate_fields(data):
    result = {}
    contract_amount = float(data["contract_amount"])
    npr_percent = int(data["npr"])
    npr_amount = round(contract_amount * npr_percent / 100, 2)
    taxable_income = round(contract_amount - npr_amount, 2)
    result["CONTRACT_AMOUNT"] = contract_amount
    result["RECOGNIZED_EXPENSES"] = npr_amount
    result["TAXABLE_INCOME"] = taxable_income

    if data["has_disability"]:
        taxable_for_tax = max(0.00, taxable_income - 7920.00)
    else:
        taxable_for_tax = taxable_income

    if data.get("manual_taxable_for_tax"):
        taxable_for_tax = float(data.get("manual_taxable_for_tax_amount", 0))

    result["TAXABLE_FOR_TAX"] = taxable_for_tax

    if data["manual_income"]:
        insurance_income = float(data["manual_income_amount"])
    elif data["npr"] == "10":
        insurance_income = 0.00
    elif taxable_income < 1077 and not data["insured_elsewhere"]:
        insurance_income = 0.00
    elif data["max_insured"]:
        insurance_income = 0.00
    else:
        income_gap = MAX_INSURANCE_INCOME - float(data["monthly_other_income"])
        insurance_income = min(taxable_income, income_gap)

    result["INSURANCE_INCOME"] = round(insurance_income, 2)

    pension = dzpo = health = 0.00
    if data["retired"] and not data["retired_wants_insurance"]:
        health = round(insurance_income * 0.032, 2)
    else:
        if data["born_after_1959"]:
            pension = round(insurance_income * 0.0658, 2)
            dzpo = round(insurance_income * 0.022, 2)
        else:
            pension = round(insurance_income * 0.0878, 2)
        health = round(insurance_income * 0.032, 2)

    result["PENSION_CONTRIBUTION"] = pension
    result["DZPO_CONTRIBUTION"] = dzpo
    result["HEALTH_CONTRIBUTION"] = health

    total_contributions = pension + dzpo + health
    result["TAXABLE_TOTAL"] = round(taxable_for_tax - total_contributions, 2)

    if taxable_for_tax == 0.00:
        result["TAXABLE_TOTAL"] = 0.00
        result["TAX_ADVANCE"] = 0.00
        result["NET_AMOUNT"] = contract_amount - total_contributions
        return result

    deduct_tax = not (data["retired"] and data["no_tax_iv_trim"] and data["doc_date"].month in [10, 11, 12])
    tax_advance = round(result["TAXABLE_TOTAL"] * 0.10, 2) if deduct_tax else 0.00
    result["TAX_ADVANCE"] = tax_advance
    result["NET_AMOUNT"] = round(contract_amount - total_contributions - tax_advance, 2)
    return result

def show_result_summary(result):
    labels = {
        "CONTRACT_AMOUNT": "1. Сума по договора",
        "RECOGNIZED_EXPENSES": "2. Признати разходи",
        "TAXABLE_INCOME": "3. Облагаем доход",
        "TAXABLE_FOR_TAX": "4. Облагаема част (ред 4)",
        "INSURANCE_INCOME": "5. Осигурителен доход (ред 5)",
        "PENSION_CONTRIBUTION": "6.1 Фонд Пенсии",
        "DZPO_CONTRIBUTION": "6.2 ДЗПО",
        "HEALTH_CONTRIBUTION": "6.3 Здравно осигуряване",
        "TAXABLE_TOTAL": "7. Сума за авансово облагане (ред 7)",
        "TAX_ADVANCE": "8. Авансов данък (ред 8)",
        "NET_AMOUNT": "9. Сума за получаване (ред 9)"
    }
    st.markdown('<div class="result-box">', unsafe_allow_html=True)
    for key in labels:
        if key in result:
            st.markdown(f"**{labels[key]}:** {result[key]:.2f}")
    st.markdown('</div>', unsafe_allow_html=True)

def replace_markers_in_paragraphs(paragraphs, data_dict):
    for para in paragraphs:
        full_text = para.text
        for key, value in data_dict.items():
            full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
        if para.text != full_text:
            for run in para.runs:
                run.text = ""
            para.add_run(full_text)


def num_to_words_bg(amount):
    units = ["", "един", "два", "три", "четири", "пет", "шест", "седем", "осем", "девет"]
    teens = ["десет", "единадесет", "дванадесет", "тринадесет", "четиринадесет", "петнадесет",
             "шестнадесет", "седемнадесет", "осемнадесет", "деветнадесет"]
    tens = ["", "", "двадесет", "тридесет", "четиридесет", "петдесет", "шестдесет",
            "седемдесет", "осемдесет", "деветдесет"]
    hundreds = ["", "сто", "двеста", "триста", "четиристотин", "петстотин",
                "шестстотин", "седемстотин", "осемстотин", "деветстотин"]

    def under_thousand(n):
        result = []
        h, rem = divmod(n, 100)
        if h < len(hundreds):
            result.append(hundreds[h])
        if 10 <= rem < 20:
            result.append(teens[rem - 10])
        else:
            t, u = divmod(rem, 10)
            if t:
                result.append(tens[t])
            if u:
                result.append(units[u])
        return " ".join(result).strip()

    leva = int(math.floor(amount))
    stotinki = int(round((amount - leva) * 100))

    if leva == 0:
        leva_words = "нула лева"
    elif leva < 1000:
        leva_words = under_thousand(leva) + " лева"
    else:
        thousands = leva // 1000
        below = leva % 1000
        if thousands == 1:
            leva_words = "хиляда"
        else:
            leva_words = under_thousand(thousands) + " хиляди"
        if below > 0:
            leva_words += " " + under_thousand(below)
        leva_words += " лева"

    if stotinki == 0:
        return leva_words
    elif stotinki == 1:
        return leva_words + " и 1 стотинка"
    else:
        return leva_words + f" и {stotinki} стотинки"



def format_quarter_checkboxes(month):
    quarters = {
        1: "☑ І-во тр. ☐ ІІ-ро тр. ☐ ІІІ-то тр. ☐ ІV-то тр.",
        2: "☐ І-во тр. ☑ ІІ-ро тр. ☐ ІІІ-то тр. ☐ ІV-то тр.",
        3: "☐ І-во тр. ☐ ІІ-ро тр. ☑ ІІІ-то тр. ☐ ІV-то тр.",
        4: "☐ І-во тр. ☐ ІІ-ро тр. ☐ ІІІ-то тр. ☑ ІV-то тр."
    }
    if month in [1,2,3]: return quarters[1]
    if month in [4,5,6]: return quarters[2]
    if month in [7,8,9]: return quarters[3]
    return quarters[4]


def render_docx_form(result_fields):
    st.subheader("Въведете данни за печат")
    company_name = st.text_input("Име на предприятието")
    company_eik = st.text_input("ЕИК")
    nap_office = st.text_input("ТД на НАП")
    person_name = st.text_input("Три имена на лицето")
    person_egn = st.text_input("ЕГН")
    contract_number = st.text_input("Номер на договора")
    contract_date = st.date_input("Дата на договора", value=date.today())
    quarter = st.selectbox("Тримесечие", ["І-во тр.", "ІІ-ро тр.", "ІІІ-то тр.", "ІV-то тр."])
    doc_date = st.session_state.get("doc_date", date.today())
    def checkbox_marked(val): return "☑ да    ☐ не" if val else "☐ да    ☑ не"
    net_amount_words = num_to_words_bg(result_fields.get("NET_AMOUNT", 0))
    filled_data = {
        "COMPANY_NAME": company_name,
        "COMPANY_EIK": company_eik,
        "NAP_OFFICE": nap_office,
        "PERSON_NAME": person_name,
        "PERSON_EGN": person_egn,
        "CONTRACT_NUMBER": contract_number,
        "CONTRACT_DATE": contract_date.strftime("%d.%m.%Y"),
        "QUARTER": quarter,
        "HAS_DISABILITY": checkbox_marked(st.session_state.get("has_disability", False)),
        "WANTS_TAX_IV_TRIM": checkbox_marked(not st.session_state.get("no_tax_iv_trim", False)) if doc_date.month in [10,11,12] else "☐ да    ☐ не",
        "MAX_INSURED": checkbox_marked(st.session_state.get("max_insured", False)),
        "RETIRED": checkbox_marked(st.session_state.get("retired", False)),
        "WANTS_INSURANCE": checkbox_marked(st.session_state.get("retired_wants_insurance", False)) if st.session_state.get("retired", False) else "",
        "INSURED_ELSEWHERE": checkbox_marked(st.session_state.get("insured_elsewhere", False)),
        "NET_AMOUNT_WORDS": net_amount_words,
        "QUARTER_CHECKBOXES": format_quarter_checkboxes(doc_date.month),
        "INSURANCE_TOTAL": f"{result_fields.get('PENSION_CONTRIBUTION', 0) + result_fields.get('DZPO_CONTRIBUTION', 0) + result_fields.get('HEALTH_CONTRIBUTION', 0):.2f}",
        "MONTH_AND_YEAR": f"{doc_date.strftime('%m.%Y')}"
    }
    if result_fields:
        filled_data.update({k: f"{v:.2f}" if isinstance(v, float) else str(v) for k, v in result_fields.items()})
    if st.button("Генерирай и изтегли бланка"):
        doc = Document("template.docx")
        replace_markers_in_paragraphs(doc.paragraphs, filled_data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_markers_in_paragraphs(cell.paragraphs, filled_data)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="⬇️ Изтегли попълнената бланка",
            data=buffer,
            file_name=f"smetka_{person_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def main():
    tab1, tab2 = st.tabs(["Калкулатор", "Печат на бланка"])
    with tab1:
        st.header("Калкулатор")
        contract_amount = st.number_input("Сума по договора (ред 1)", value=0.00)
        npr = st.selectbox("Нормативно признати разходи (%)", options=["10", "25", "40", "60"], index=1)
        has_disability = st.checkbox("Лице с увреждане ≥ 50%")
        max_insured = st.checkbox("Осигурен върху максималния осигурителен доход")
        retired = st.checkbox("Пенсионер съм")
        retired_wants_insurance = False
        if retired:
            retired_wants_insurance = st.checkbox("Желая да бъда осигуряван за фонд 'Пенсии' и ДЗПО")
        insured_elsewhere = st.checkbox("Осигурен съм на друго основание през месеца")
        monthly_other_income = st.number_input("Месечен доход без тази сума (ако не е на макс)", value=0.00)
        born_after_1959 = st.checkbox("Роден съм след 31.12.1959 г.")
        manual_income = st.checkbox("Ръчно въведи ред 5")
        manual_income_amount = st.number_input("Ред 5 - Осигурителен доход", value=0.00) if manual_income else 0.00
        manual_taxable_for_tax = st.checkbox("Ръчно въведи ред 4 - Облагаема част")
        manual_taxable_for_tax_amount = st.number_input("Ред 4 - Облагаема част", value=0.00) if manual_taxable_for_tax else 0.00
        doc_date = st.date_input("Дата на сметката", value=date.today())
        no_tax_iv_trim = st.checkbox("Не желая удържане на данък (IV тримесечие)") if doc_date.month in [10, 11, 12] else False

        data = {
            "contract_amount": contract_amount,
            "npr": npr,
            "has_disability": has_disability,
            "no_tax_iv_trim": no_tax_iv_trim,
            "max_insured": max_insured,
            "retired": retired,
            "retired_wants_insurance": retired_wants_insurance,
            "insured_elsewhere": insured_elsewhere,
            "monthly_other_income": monthly_other_income,
            "born_after_1959": born_after_1959,
            "manual_income": manual_income,
            "manual_income_amount": manual_income_amount,
            "manual_taxable_for_tax": manual_taxable_for_tax,
            "manual_taxable_for_tax_amount": manual_taxable_for_tax_amount,
            "doc_date": doc_date,
        }

        result = None
        if st.button("Изчисли"):
            result = calculate_fields(data)
            st.session_state["last_result"] = result
            st.session_state["has_disability"] = has_disability
            st.session_state["no_tax_iv_trim"] = no_tax_iv_trim
            st.session_state["max_insured"] = max_insured
            st.session_state["retired"] = retired
            st.session_state["retired_wants_insurance"] = retired_wants_insurance
            st.session_state["insured_elsewhere"] = insured_elsewhere
            st.session_state["doc_date"] = doc_date
            show_result_summary(result)

    with tab2:
        st.header("Печат на бланка")
        if "last_result" in st.session_state and st.session_state["last_result"]:
            render_docx_form(st.session_state["last_result"])
        else:
            st.info("Моля, първо изчислете сумите в таба 'Калкулатор'.")

if __name__ == "__main__":
    main()
